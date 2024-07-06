"""
This is RedLine Agent
This module contains the document comparision logic.

Input:
Take two inputs, a gold standard text and a text to be reviewed
The gold standard text is indexed in Azure AI Search
The review text is formatted as JSON that retruned from Document AI Service

Output:
Output of this module is a Word that has commnets generated by the AI 

Logic:
Two level of comparisions are included:
1. Every 400 words (Text) include 10% of overlaped words 
1.1. Check section number and title
1.2. Calculate similarity of given text with the gold standard text

2. Each paragraph (Microsoft Words)
"""

# RedLine Agent class

import os
import numpy as np
from docx import Document
from dotenv import load_dotenv
from openai import AzureOpenAI
from azure.storage.blob import BlobServiceClient
from azure.core.exceptions import ResourceExistsError
from azure.ai.documentintelligence.models import AnalyzeResult 
from langchain_community.vectorstores.azuresearch import AzureSearch
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import AnalyzeDocumentRequest, ContentFormat, AnalyzeResult 
import logging  
  

class RedLineAgent:
    def __init__(self, 
                 openai: AzureOpenAI, 
                 gold_search: AzureSearch, 
                 source_file_full_path:str, 
                 container_name: str = "ada-container"):
        # # Configure root logger to ignore DEBUG and INFO messages  
        logging.basicConfig(level=logging.WARNING)  
        logging.getLogger('azure').setLevel(logging.WARNING)  
        
        env_file_path = '.env'
        load_dotenv(env_file_path, override=True)

        assert isinstance(openai, AzureOpenAI), "OpenAI should be an instance of AsyncAzureOpenAI"
        assert isinstance(gold_search, AzureSearch), "Gold search should be an instance of AzureSearch"
        assert isinstance(source_file_full_path, str), "File path should be a string"

        self.openai = openai
        self.gold_search = gold_search
        self.blob_service_client = BlobServiceClient.from_connection_string(os.getenv('AZURE_STORAGE_CONNECTION_STRING'))
        self.source_file_full_path = source_file_full_path
        self.container_name = container_name


    def run(self):
        blob_client = self.upload_document()

        document_intelligent_result = self.run_document_analysis(blob_client)
        document_intelligent_contents = [paragraph.content for paragraph in document_intelligent_result.paragraphs]

        examiner_comments = self.examine(document_intelligent_contents)
        blob, reviewed_docx_path = self.add_save_examiner_comment(examiner_comments, blob_client)

        return blob, reviewed_docx_path
        

    def upload_document(self):
                # check container exsitence
        # get file name from source_file_full_path
        file_name = os.path.basename(self.source_file_full_path)
        
        try:
            self.blob_service_client.create_container(self.container_name)
        except ResourceExistsError:
            pass

        blob_client = self.blob_service_client.get_blob_client(container=self.container_name, blob=file_name)

        with open(self.source_file_full_path, 'rb') as data:
            blob_client.upload_blob(data, overwrite=True)

        return blob_client
    
    def examine(self, examinee_paragraphs, top_k=1):
        """
        Use Azure OpenAI to compare the examinee paragraphs with the gold standard
        """

        results = []

        for i, examinee_paragraph in enumerate(examinee_paragraphs):
            # get the gold standard paragraph
            docs = docs = self.gold_search.similarity_search(
                query=examinee_paragraph,
                k=top_k,
                search_type='hybrid'
            )

            # gold_paragraph = [doc.page_content for doc in docs]
            gold_paragraph = docs[0].page_content
            # compare the two paragraphs
            # get the similarity score
            similarity_score = self.similarity(str(gold_paragraph), str(examinee_paragraph))
            examiner_comment = self.generate_comment(similarity_score, gold_paragraph, examinee_paragraph)
            results.append({
                "gold_paragraph": gold_paragraph,
                "paragraph": examinee_paragraph,
                "similarity_score": similarity_score,
                "examiner_comment": examiner_comment
            })
        return results

    def generate_comment(self, similarity_score, gold_paragraph, examinee_paragraph):
        system_prompt = """
            You are a contract reviewer in a business negotiation team. 
            You are responsible for reviewing contracts and providing insights. 

            ## Review process
            There are two versions of the paragraphs with cosine similarity. 
            If the cosine similarity is less than 0.9, then the human reviewer needs to review the differences.
            gold paragraph: The gold standard in our company
            examinee paragraph: Modified paragraph by the client
            The paragraphs are similar and have overlapping sentences.
            Find overlapping sentences between the two documents. 
            And compare two document from the overlapping sentences.

            ## Examination guideline
            ### 1. Response format rule
            Response in simple text format that wraped with '```'
            you can add additional information about the difference. added, removed, or modified.

            ### 2. When the two paragraphs are identical
            Use 'No differences' in the comment field. 

            ### 3. When the two paragraphs are off the topic or completely different
            Use "Not able to compare" in the comment field.

            ### 4. When the two paragraphs are similar but not identical
            Return the description of the difference in the comment field
            If there is potential risk for our company, or against the gold standard, please provide the details in 20 words

            ### 5. Out of scope
            If examinee paragraph is not related to the gold paragraph, please provide the details in 20 words
            Here is scope of gold standard, if the given paragraph is not related to the gold paragraph, it is out of scope.
            #### Scope
            1.1 Acceptance
            1.2 Acceptance Criteria
            1.3 Acceptance Test
            1.4 Buyer Spend
            1.5 Confidential Information
            1.6 Consumable Item(s)
            1.7 Counterfeit(s)
            1.8 Custom Item(s)
            1.9 Deliver, Delivered or Delivery
            1.10 Delivery Date
            1.11 Delivery Point
            1.12 Demand Flow Technology (“DFT”) Signal Alert
            1.13 End of Life (“EOL”) Date
            1.14 Engineering Change Order (“ECO”)
            1.15 Epidemic Failure
            1.16 Field
            1.17 Forecast(s)
            1.18 Hazardous Materials
            1.19 Intellectual Property (“IP”) Rights
            1.20 Items
            1.21 Inventions
            1.22 Lead-time
            1.23 Margin
            1.24 Mark-Up
            1.25 Milestone
            1.26 Milestone Commitment Date
            1.27 P1 Request
            1.28 P2 Request
            1.29 Project
            1.30 Purchase Order
            1.31 Purchase or Product Specifications
            1.32 Release
            1.33 Schedule(s)
            1.34 Service(s)
            1.35 Standard Item
            1.36 Statement of Work
            1.37 Source Inspection
            1.38 Technology Roadmap
            2.1 Purchase Orders
            2.2 Schedules
            2.3 Changes and Cancellations
            2.4 Delivery Terms
            2.5 Shipment and Packaging
            2.6 Inspection and Acceptance
            2.7 Transfer of Title and Risk of Loss
            2.8 Source Inspection

            ## Examples
            ### Example When the two paragraphs are similar but not identical
            The cosine similarity between the paragraph is 0.993.

            ---gold---
            Access control can be implemented at different levels of the database, such as the schema, table, column, row, or cell level.
            ------

            ---examinee---
            Access control can be implemented at different levels of the database, such as the schema, table, column, row, or cell level. The view and the stored procedures also need to have controlled access.
            ------

            Your Answer:
            ```
            More conditions are added to the access control. There is no risk for the company.
            ```

            ### Example When the two versions are identical
            The cosine similarity between the paragraph is 0.992.

            ---gold---
            Access control can be implemented at different levels of the database.
            ------

            ---examinee---
            Access control can be implemented at different levels of the database.
            ------

            Your Answer:
            ```
            No differences
            ```

            ### Example When the two versions are off the topic or completely different
            The cosine similarity between the paragraph is 0.990.

            ---gold---
            2.4 Delivery Terms. Contoso shall Deliver the Items or the Services to QBIT Maker
            ------

            ---examinee---
            Americano is made with espresso and hot water.
            ------

            Your Answer:
            ```
            Not able to compare
            ```

            ### Example out of socpe
            The cosine similarity between the paragraph is 0.123.

            ---gold---
            
            ------

            ---examinee---
            3.0 Blueprint design
            ------

            Your Answer:
            ```
            out of scope
            ```
            """
        
        user_prompt = """ 
            Compare the two paragraphs and provide the difference in the field
            The cosine similarity between the paragraph is {{similarity_score}}

            ---gold---
            {{gold_paragraph}}
            ------

            ---examinee---
            {{examinee_paragraph}}
            ------
            """
        
        updated_user_prompt = user_prompt.replace("{{gold_paragraph}}", gold_paragraph).\
                                        replace("{{examinee_paragraph}}", examinee_paragraph).\
                                        replace("{{similarity_score}}", str(similarity_score))
                                                                                                
        res = self.openai.chat.completions.create(
            model = os.getenv("DEPLOYMENT_NAME"),
            messages = [
                {
                    "role": "system",
                    "content": system_prompt
                },
                {
                    "role": "user",
                    "content": updated_user_prompt
                }
            ],
            temperature=0.1
        )
        
        return res.choices[0].message.content.split("```")[1].strip()

    def add_save_examiner_comment(self, examiner_comments, blob_client):
        temp_dir_path = 'temp'
        os.makedirs(temp_dir_path, exist_ok=True)
        examinee_docx_path = os.path.join(temp_dir_path, blob_client.blob_name)
        
        # download the document
        with open(examinee_docx_path, 'wb') as data:
            blob_client.download_blob().readinto(data)        

        # read the document
        examinee_docx = Document(examinee_docx_path)
        examinee_paragraphs = [p.text for p in examinee_docx.paragraphs]

        # check if the number of paragraphs in the examinee document and the number of examiner comments match
        if len(examinee_paragraphs) != len(examiner_comments) and  examinee_paragraphs[-1] == "":
            #raise ValueError("The number of paragraphs in the examinee document and the number of examiner comments do not match")
            examinee_paragraphs.pop(-1)
                        
        # print(f"Len of examinee paragraphs: {len(examinee_paragraphs)}\nLen of examiner comments: {len(examiner_comments)}")        

        for i, paragraph in enumerate(examinee_paragraphs):
            if i <= len(examiner_comments)-1:
                if examiner_comments[i]["examiner_comment"] != "No differences":
                    examinee_docx.paragraphs[i].add_comment(f"{examiner_comments[i]['examiner_comment']}", "AI Examiner")
                else:
                    pass
        # Update file name adding 'redline' prefix in file name
        reviewed_docx_path = examinee_docx_path.replace(".docx", "_redline.docx")
        examinee_docx.save(reviewed_docx_path)
        
        upload_blob_client = self.blob_service_client.get_blob_client(container=self.container_name, blob=os.path.basename(reviewed_docx_path))
        with open(reviewed_docx_path, 'rb') as data:
            upload_blob_client.upload_blob(data, overwrite=True)

        return upload_blob_client, reviewed_docx_path




    def similarity(self, gold_paragraph:str, examinee:str, engine: str=os.getenv("EMBEDDING_MODEL_NAME")):
        if len(gold_paragraph) <= 1 or len(examinee) <= 1:
            return 0
        
        res = self.openai.embeddings.create(
            model=engine,
            input=[gold_paragraph, examinee],
        )

        t1, t2 = res.data[0].embedding, res.data[1].embedding
        
        return np.dot(t1, t2) / (np.linalg.norm(t1) * np.linalg.norm(t2)) 

    def run_document_analysis(self, blob_client):
        """
        Get the document analysis
        """
        # get the document analysis from blob storage
        # log the document analysis in table
        endpoint = os.getenv('AZURE_FORM_RECOGNIZER_ENDPOINT')
        key = os.getenv('AZURE_FORM_RECOGNIZER_KEY')

        document_intelligence_client  = DocumentIntelligenceClient(endpoint=endpoint, credential=AzureKeyCredential(key))

        document_bytes = blob_client.download_blob().readall()

        poller = document_intelligence_client.begin_analyze_document(model_id="prebuilt-layout",
                                                                    analyze_request=AnalyzeDocumentRequest(bytes_source=document_bytes),
                                                                    output_content_format=ContentFormat.TEXT)
        
        result: AnalyzeResult = poller.result()

        return result